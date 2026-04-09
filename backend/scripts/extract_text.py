#!/usr/bin/env python3
"""
Script para extraer texto de archivos PDF y DOCX.
Imprime JSON estructurado al stdout.

Usage: python extract_text.py <ruta_archivo>
"""

import json
import sys
import os
from pathlib import Path

# Intentar importar pdfplumber y python-docx
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None


def extract_from_pdf(file_path: str) -> dict:
    """Extrae texto de un archivo PDF usando pdfplumber."""
    if pdfplumber is None:
        return {
            "success": False,
            "error": "Biblioteca pdfplumber no está instalada.",
            "code": "MISSING_DEPENDENCY",
        }

    if not os.path.exists(file_path):
        return {
            "success": False,
            "error": f"El archivo no existe: {file_path}",
            "code": "FILE_NOT_FOUND",
        }

    try:
        with pdfplumber.open(file_path) as pdf:
            num_pages = len(pdf.pages)

            # Extraer texto de todas las páginas
            full_text = ""
            has_content = False

            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    has_content = True
                full_text += page_text + "\n"

            full_text = full_text.strip()

            # Si no hay contenido extraíble, es probable que sea un PDF escaneado
            if not has_content or not full_text:
                return {
                    "success": False,
                    "error": "El archivo no contiene texto extraíble. Usa un PDF con texto seleccionable.",
                    "code": "NO_TEXT",
                }

            return {
                "success": True,
                "text": full_text,
                "filename": os.path.basename(file_path),
                "pages": num_pages,
            }

    except Exception as e:
        return {
            "success": False,
            "error": f"Error al procesar PDF: {str(e)}",
            "code": "PDF_ERROR",
        }


def extract_from_docx(file_path: str) -> dict:
    """Extrae texto de un archivo DOCX usando python-docx."""
    if docx is None:
        return {
            "success": False,
            "error": "Biblioteca python-docx no está instalada.",
            "code": "MISSING_DEPENDENCY",
        }

    if not os.path.exists(file_path):
        return {
            "success": False,
            "error": f"El archivo no existe: {file_path}",
            "code": "FILE_NOT_FOUND",
        }

    try:
        doc = docx.Document(file_path)

        # Extraer texto de todos los párrafos
        full_text = "\n".join([para.text for para in doc.paragraphs])

        # También extraer texto de tablas si las hay
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text:
                    full_text += "\n" + row_text

        full_text = full_text.strip()

        if not full_text:
            return {
                "success": False,
                "error": "El archivo DOCX está vacío.",
                "code": "NO_TEXT",
            }

        # python-docx no proporciona número de páginas directamente,
        # podemos contar las secciones o simplemente indicar 1
        num_sections = len(doc.sections)

        return {
            "success": True,
            "text": full_text,
            "filename": os.path.basename(file_path),
            "pages": num_sections,  # Usamos secciones como aproximación
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Error al procesar DOCX: {str(e)}",
            "code": "DOCX_ERROR",
        }


def main():
    if len(sys.argv) < 2:
        print(
            json.dumps(
                {
                    "success": False,
                    "error": "Uso: python extract_text.py <ruta_archivo>",
                    "code": "INVALID_ARGS",
                }
            )
        )
        sys.exit(1)

    file_path = sys.argv[1]

    # Determinar el tipo de archivo por extensión
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        result = extract_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        # Para .docx usamos python-docx
        result = extract_from_docx(file_path)
    else:
        result = {
            "success": False,
            "error": f"Tipo de archivo no soportado: {ext}. Solo se aceptan PDF, DOC y DOCX.",
            "code": "UNSUPPORTED_TYPE",
        }

    # Imprimir resultado como JSON
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
